import sys
import json
from docx import Document

def preencher_ficha(template_path, output_path, dados):
    doc = Document(template_path)

    # Preenche os campos na primeira ficha
    for paragraph in doc.paragraphs:
        text = paragraph.text

        if 'PROFESSOR:' in text and dados.get('professor'):
            for run in paragraph.runs:
                if 'PROFESSOR:' in run.text:
                    run.text = f'PROFESSOR: {dados["professor"]}'

        if 'COMPONENTE/DISCIPLINA:' in text and dados.get('disciplina'):
            for run in paragraph.runs:
                if 'COMPONENTE/DISCIPLINA:' in run.text:
                    run.text = f'COMPONENTE/DISCIPLINA: {dados["disciplina"]}'

        if 'ESTUDANTE:' in text and dados.get('aluno'):
            for run in paragraph.runs:
                if 'ESTUDANTE:' in run.text:
                    run.text = f'ESTUDANTE: {dados["aluno"]}'

        if 'TURMA:' in text and dados.get('turma'):
            for run in paragraph.runs:
                if 'TURMA:' in run.text:
                    run.text = f'TURMA: {dados["turma"]}'

        if 'PEDAGOGA RESP.' in text and dados.get('pedagoga'):
            for run in paragraph.runs:
                if 'PEDAGOGA RESP.' in run.text:
                    run.text = f'PEDAGOGA RESP.: {dados["pedagoga"]}'

        if 'QUINZENA/DATA:' in text and dados.get('data'):
            for run in paragraph.runs:
                if 'QUINZENA/DATA:' in run.text:
                    run.text = f'QUINZENA/DATA: {dados["data"]}'

        if 'Nº DE AULAS:' in text and dados.get('numAulas'):
            for run in paragraph.runs:
                if 'Nº DE AULAS:' in run.text:
                    run.text = f'Nº DE AULAS: {dados["numAulas"]}'

        if 'ENCAMINHAMENTO:' in text and dados.get('encaminhamento'):
            for run in paragraph.runs:
                if 'ENCAMINHAMENTO:' in run.text:
                    run.text = f'ENCAMINHAMENTO: {dados["encaminhamento"]}'

        if 'ROTEIRO DE ESTUDOS' in text and dados.get('roteiro'):
            for run in paragraph.runs:
                if 'ROTEIRO DE ESTUDOS' in run.text:
                    run.text = f'ROTEIRO DE ESTUDOS: {dados["roteiro"]}'

        if 'OBSERVAÇÕES:' in text and dados.get('observacoes'):
            for run in paragraph.runs:
                if 'OBSERVAÇÕES:' in run.text:
                    run.text = f'OBSERVAÇÕES: {dados["observacoes"]}'

    # Preenche tabelas se houver
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if 'PROFESSOR:' in text and dados.get('professor'):
                        for run in paragraph.runs:
                            if 'PROFESSOR:' in run.text:
                                run.text = f'PROFESSOR: {dados["professor"]}'
                    if 'DISCIPLINA:' in text and dados.get('disciplina'):
                        for run in paragraph.runs:
                            if 'DISCIPLINA:' in run.text:
                                run.text = f'DISCIPLINA: {dados["disciplina"]}'

    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    dados_json = sys.argv[3]
    dados = json.loads(dados_json)

    result = preencher_ficha(template_path, output_path, dados)
    print(result)
